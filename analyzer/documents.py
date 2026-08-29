from io import BytesIO

from django.http import HttpResponse

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
)

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def clean_lines(text):
    if not text:
        return []

    return [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]


def create_pdf_response(resume):

    buffer = BytesIO()

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )

    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "ResumeName",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=5,
    )

    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        alignment=TA_CENTER,
        spaceAfter=5,
    )

    contact_style = ParagraphStyle(
        "ResumeContact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11,
        alignment=TA_CENTER,
        spaceAfter=10,
    )

    section_style = ParagraphStyle(
        "ResumeSection",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=13,
        spaceBefore=8,
        spaceAfter=5,
    )

    content_style = ParagraphStyle(
        "ResumeContent",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        spaceAfter=2,
    )

    story = []

    # Name
    story.append(
        Paragraph(
            resume.full_name,
            name_style
        )
    )

    # Professional title
    if resume.professional_title:
        story.append(
            Paragraph(
                resume.professional_title,
                title_style
            )
        )

    # Contact details
    contact = []

    if resume.email:
        contact.append(resume.email)

    if resume.phone:
        contact.append(resume.phone)

    if resume.location:
        contact.append(resume.location)

    if resume.linkedin:
        contact.append(resume.linkedin)

    if resume.github:
        contact.append(resume.github)

    if contact:
        story.append(
            Paragraph(
                " | ".join(contact),
                contact_style
            )
        )

    story.append(
        HRFlowable(
            width="100%",
            thickness=0.7,
            spaceAfter=7,
        )
    )

    def add_section(title, content):

        if not content:
            return

        story.append(
            Paragraph(
                title.upper(),
                section_style
            )
        )

        for line in clean_lines(content):

            escaped = (
                line
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )

            if escaped.startswith("•"):
                escaped = "&bull; " + escaped[1:].strip()

            elif escaped.startswith("-"):
                escaped = "&bull; " + escaped[1:].strip()

            story.append(
                Paragraph(
                    escaped,
                    content_style
                )
            )

    add_section(
        "Professional Summary",
        resume.summary
    )

    add_section(
        "Technical Skills",
        resume.skills
    )

    add_section(
        "Experience",
        resume.experience
    )

    add_section(
        "Projects",
        resume.projects
    )

    add_section(
        "Education",
        resume.education
    )

    add_section(
        "Certifications",
        resume.certifications
    )

    add_section(
        "Achievements",
        resume.achievements
    )

    document.build(story)

    buffer.seek(0)

    filename = (
        f"{resume.full_name}_Resume.pdf"
        .replace(" ", "_")
    )

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/pdf",
    )

    response["Content-Disposition"] = (
        f'attachment; filename="{filename}"'
    )

    return response


def create_docx_response(resume):

    document = Document()

    section = document.sections[0]

    section.top_margin = 0.6
    section.bottom_margin = 0.6
    section.left_margin = 0.7
    section.right_margin = 0.7

    normal_style = document.styles["Normal"]

    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)

    # Name
    paragraph = document.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        resume.full_name
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)

    # Professional title
    if resume.professional_title:

        paragraph = document.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            resume.professional_title
        )

        run.font.name = "Arial"
        run.font.size = Pt(10)

    # Contact
    contact = []

    if resume.email:
        contact.append(resume.email)

    if resume.phone:
        contact.append(resume.phone)

    if resume.location:
        contact.append(resume.location)

    if resume.linkedin:
        contact.append(resume.linkedin)

    if resume.github:
        contact.append(resume.github)

    if contact:

        paragraph = document.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            " | ".join(contact)
        )

        run.font.name = "Arial"
        run.font.size = Pt(8.5)

    def add_section(title, content):

        if not content:
            return

        paragraph = document.add_paragraph()

        run = paragraph.add_run(
            title.upper()
        )

        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)

        for line in clean_lines(content):

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(2)

            run = paragraph.add_run(line)

            run.font.name = "Arial"
            run.font.size = Pt(9.5)

    add_section(
        "Professional Summary",
        resume.summary
    )

    add_section(
        "Technical Skills",
        resume.skills
    )

    add_section(
        "Experience",
        resume.experience
    )

    add_section(
        "Projects",
        resume.projects
    )

    add_section(
        "Education",
        resume.education
    )

    add_section(
        "Certifications",
        resume.certifications
    )

    add_section(
        "Achievements",
        resume.achievements
    )

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    filename = (
        f"{resume.full_name}_Resume.docx"
        .replace(" ", "_")
    )

    response = HttpResponse(
        buffer.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    response["Content-Disposition"] = (
        f'attachment; filename="{filename}"'
    )

    return response